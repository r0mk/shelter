#!/usr/bin/python3
#bot = telegram.Bot(token='255961491:AAHsOehoD9TUUH7aakSGZGRPqxy2jSmwBn4')

from telegram import InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Updater,  CallbackQueryHandler
updater = Updater(token='255961491:AAHsOehoD9TUUH7aakSGZGRPqxy2jSmwBn4')
dispatcher = updater.dispatcher
import logging
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
import time
from datetime import date
import docx

def duty(bot,update):
    today=date.fromtimestamp(time.time()).day
    worddoc = docx.Document ("duty.docx")
    today+=0
    duty_order=[]
    for row in worddoc.tables[1].rows:
        duty_order_row = []
        for cell in row.cells:
            duty_order_row.append(cell.text)
        duty_order.append(duty_order_row)
    
    if int(time.strftime("%H", time.gmtime())) < 8:
        today+=0
    else:
        today+=1

    result = []
    for i in duty_order:
        if i[today] == 'x':
            result.append(i[1])
    text = 'Дежурят до 8 утра ' + ', '.join(set(result))
    
    bot.sendMessage(chat_id=update.message.chat_id, text=text)


def duty_tomorrow(bot,update):
    today=date.fromtimestamp(time.time()).day
    worddoc = docx.Document ("duty.docx")
    today+=0
    duty_order=[]
    for row in worddoc.tables[1].rows:
        duty_order_row = []
        for cell in row.cells:
            duty_order_row.append(cell.text)
        duty_order.append(duty_order_row)

    if int(time.strftime("%H", time.gmtime())) < 8:
        today+=1
    else:
        today+=2

    result = []
    for i in duty_order:
        if i[today] == 'x':
           result.append(i[1])
    text = 'Будут дежурить c 8 утра ' + ', '.join(set(result))
    bot.sendMessage(chat_id=update.message.chat_id, text=text)

def start(bot, update):
    bot.sendMessage(chat_id=update.message.chat_id, text="/duty дежурные\n /duty_tomorrow - держурят завтра \n /send_schedule график дежурств на месяц")

def help(bot, update):
    update.message.reply_text("/duty дежурные\n /duty_tomorrow - держурят завтра \n /send_schedule график дежурств на месяц")

def send_schedule(bot, update):
    bot.sendDocument(chat_id=update.message.chat_id, document=open('duty.docx', 'rb'))


from telegram.ext import CommandHandler
help_handler = CommandHandler('help', help)
dispatcher.add_handler(help_handler)

duty_handler = CommandHandler('duty', duty)
dispatcher.add_handler(duty_handler)

send_schedule_handler = CommandHandler('send_schedule', send_schedule)
dispatcher.add_handler(send_schedule_handler)

duty_tomorrow_handler = CommandHandler('duty_tomorrow', duty_tomorrow)
dispatcher.add_handler(duty_tomorrow_handler)

start_handler = CommandHandler('start', start)
dispatcher.add_handler(start_handler)
updater.start_polling()

def echo(bot, update):
    bot.sendMessage(chat_id=update.message.chat_id, text=update.message.text)

from telegram.ext import MessageHandler, Filters
echo_handler = MessageHandler(Filters.text, echo)
dispatcher.add_handler(echo_handler)

def caps(bot, update, args):
    text_caps = ' '.join(args).upper()
    bot.sendMessage(chat_id=update.message.chat_id, text=text_caps)

caps_handler = CommandHandler('caps', caps, pass_args=True)
dispatcher.add_handler(caps_handler)

from telegram import InlineQueryResultArticle, InputTextMessageContent
def inline_caps(bot, update):
    query = update.inline_query.query
    if not query:
        return
    results = list()
    results.append(
        InlineQueryResultArticle(
            id=query.upper(),
            title='Caps',
            input_message_content=InputTextMessageContent(query.upper())
        )
    )
    bot.answerInlineQuery(update.inline_query.id, results)

from telegram.ext import InlineQueryHandler
inline_caps_handler = InlineQueryHandler(inline_caps)
dispatcher.add_handler(inline_caps_handler)


def unknown(bot, update):
    bot.sendMessage(chat_id=update.message.chat_id, text="Sorry, I didn't understand that command.")

unknown_handler = MessageHandler(Filters.command, unknown)
dispatcher.add_handler(unknown_handler)
